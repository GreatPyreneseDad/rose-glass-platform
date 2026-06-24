#!/usr/bin/env python3
"""
Leelanau County policies scraper.

Scrapes every document linked from the county policies page and consolidates
them into a single Markdown file.

Source page:
    https://www.leelanau.gov/leelanau_county/county_policies.php

What it does
------------
1. Fetches the policies index page (with a browser User-Agent; the site rejects
   the default urllib/requests agent with HTTP 403).
2. Extracts every document link (PDF / DOC / DOCX and generic file-download
   handlers) along with its visible title.
3. Downloads each document into a local cache (so re-runs are cheap and you can
   inspect the raw files).
4. Extracts the text of each document:
       - PDFs   -> PyMuPDF (fitz), falling back to pdfminer.six
       - DOCX   -> python-docx
       - DOC    -> left as a download link (binary legacy format)
5. Writes one combined Markdown file with a table of contents and one section
   per document.

Usage
-----
    pip install -r requirements-scraper.txt
    python leelanau_policies_scraper.py \
        --out leelanau_county_policies.md \
        --cache ./.policy_cache

Notes
-----
- The script is resilient: a failure on a single document is logged and the run
  continues. Failed documents are still listed in the output with the error and
  a direct link.
- Be polite: a small delay is applied between requests. Tune with --delay.
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import sys
import time
from dataclasses import dataclass, field
from urllib.parse import urljoin, urlparse, unquote

import requests
from bs4 import BeautifulSoup

POLICIES_URL = "https://www.leelanau.gov/leelanau_county/county_policies.php"

# The county web server blocks non-browser User-Agents with a 403, so present a
# realistic browser UA on every request.
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,"
        "application/pdf,*/*;q=0.8"
    ),
    "Accept-Language": "en-US,en;q=0.9",
}

# Extensions we treat as documents.
DOC_EXTENSIONS = (".pdf", ".doc", ".docx", ".rtf", ".txt", ".xls", ".xlsx")


@dataclass
class Document:
    title: str
    url: str
    ext: str = ""
    local_path: str = ""
    text: str = ""
    error: str = ""
    meta: dict = field(default_factory=dict)


# --------------------------------------------------------------------------- #
# Networking
# --------------------------------------------------------------------------- #
def make_session() -> requests.Session:
    session = requests.Session()
    session.headers.update(HEADERS)
    return session


def fetch(session: requests.Session, url: str, *, timeout: int = 45) -> requests.Response:
    resp = session.get(url, timeout=timeout, allow_redirects=True)
    resp.raise_for_status()
    return resp


# --------------------------------------------------------------------------- #
# Link extraction
# --------------------------------------------------------------------------- #
def looks_like_document(href: str) -> bool:
    """Heuristic: does this href point at a downloadable document?"""
    if not href:
        return False
    path = urlparse(href).path.lower()
    if path.endswith(DOC_EXTENSIONS):
        return True
    # Many CMS platforms serve files through a handler, e.g.
    # /files/...,  /vertical/Sites/...,  ?download=,  /DocumentCenter/View/123
    if re.search(r"/(files|documents?|documentcenter|uploads?|media)/", path):
        return True
    if "download" in urlparse(href).query.lower():
        return True
    return False


def clean_title(text: str, fallback_url: str) -> str:
    title = " ".join((text or "").split()).strip()
    if not title:
        # Derive from the file name.
        name = os.path.basename(urlparse(fallback_url).path)
        title = unquote(os.path.splitext(name)[0]).replace("_", " ").replace("-", " ").strip()
    return title or "Untitled document"


def extract_links(html: str, base_url: str) -> list[Document]:
    soup = BeautifulSoup(html, "html.parser")

    # Prefer the main content region if the template exposes one; otherwise scan
    # the whole document.
    container = (
        soup.find(id=re.compile(r"content", re.I))
        or soup.find("main")
        or soup.find("article")
        or soup
    )

    docs: list[Document] = []
    seen: set[str] = set()
    for a in container.find_all("a", href=True):
        href = a["href"].strip()
        if href.lower().startswith(("mailto:", "javascript:", "tel:", "#")):
            continue
        full = urljoin(base_url, href)
        if not looks_like_document(full):
            continue
        # Normalize to dedupe (ignore fragments).
        norm = full.split("#", 1)[0]
        if norm in seen:
            continue
        seen.add(norm)
        ext = os.path.splitext(urlparse(norm).path)[1].lower().lstrip(".")
        docs.append(Document(title=clean_title(a.get_text(), full), url=full, ext=ext))
    return docs


# --------------------------------------------------------------------------- #
# Download + text extraction
# --------------------------------------------------------------------------- #
def cache_path(cache_dir: str, doc: Document) -> str:
    name = os.path.basename(urlparse(doc.url).path) or "document"
    name = unquote(name)
    if not os.path.splitext(name)[1]:
        digest = hashlib.sha1(doc.url.encode()).hexdigest()[:8]
        ext = doc.ext or "bin"
        name = f"{name or 'document'}-{digest}.{ext}"
    return os.path.join(cache_dir, name)


def download(session: requests.Session, doc: Document, cache_dir: str) -> None:
    os.makedirs(cache_dir, exist_ok=True)
    path = cache_path(cache_dir, doc)
    if os.path.exists(path) and os.path.getsize(path) > 0:
        doc.local_path = path
        return
    resp = fetch(session, doc.url)
    with open(path, "wb") as fh:
        fh.write(resp.content)
    doc.local_path = path
    ctype = resp.headers.get("Content-Type", "")
    if ctype:
        doc.meta["content_type"] = ctype
    if not doc.ext and "pdf" in ctype:
        doc.ext = "pdf"


def extract_pdf_text(path: str) -> str:
    """Try PyMuPDF first (fast, accurate), then pdfminer.six."""
    try:
        import fitz  # PyMuPDF

        parts = []
        with fitz.open(path) as pdf:
            for page in pdf:
                parts.append(page.get_text("text"))
        text = "\n".join(parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        # Fall through to pdfminer.
        sys.stderr.write(f"  PyMuPDF failed ({exc}); trying pdfminer\n")

    try:
        from pdfminer.high_level import extract_text

        return extract_text(path) or ""
    except ImportError:
        raise RuntimeError(
            "No PDF backend available. Install PyMuPDF or pdfminer.six "
            "(see requirements-scraper.txt)."
        )


def extract_docx_text(path: str) -> str:
    from docx import Document as DocxDocument  # python-docx

    docx = DocxDocument(path)
    return "\n".join(p.text for p in docx.paragraphs)


def extract_text(doc: Document) -> None:
    ext = doc.ext or os.path.splitext(doc.local_path)[1].lstrip(".").lower()
    if ext == "pdf":
        doc.text = extract_pdf_text(doc.local_path)
    elif ext == "docx":
        doc.text = extract_docx_text(doc.local_path)
    elif ext == "txt":
        with open(doc.local_path, "r", encoding="utf-8", errors="replace") as fh:
            doc.text = fh.read()
    else:
        # .doc / .rtf / spreadsheets: don't attempt binary parsing here.
        doc.text = ""
        doc.meta["note"] = f"Binary/unsupported format (.{ext}); link preserved."


def normalize_text(text: str) -> str:
    """Tidy extracted text for Markdown: collapse runaway blank lines, strip
    trailing whitespace, drop obvious form-feed artifacts."""
    text = text.replace("\x0c", "\n\n")  # form feed -> paragraph break
    lines = [ln.rstrip() for ln in text.splitlines()]
    out: list[str] = []
    blanks = 0
    for ln in lines:
        if ln.strip() == "":
            blanks += 1
            if blanks <= 2:
                out.append("")
        else:
            blanks = 0
            out.append(ln)
    return "\n".join(out).strip()


# --------------------------------------------------------------------------- #
# Markdown assembly
# --------------------------------------------------------------------------- #
def slugify(text: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    slug = re.sub(r"[\s_-]+", "-", slug).strip("-")
    return slug or "section"


def build_markdown(docs: list[Document], source_url: str) -> str:
    today = time.strftime("%Y-%m-%d")
    lines: list[str] = []
    lines.append("# Leelanau County Policies")
    lines.append("")
    lines.append(f"_Compiled from [{source_url}]({source_url}) on {today}._")
    lines.append("")
    lines.append(f"**{len(docs)} documents.**")
    lines.append("")

    # Table of contents.
    lines.append("## Contents")
    lines.append("")
    slugs: list[str] = []
    counts: dict[str, int] = {}
    for i, doc in enumerate(docs, 1):
        base = slugify(doc.title)
        counts[base] = counts.get(base, 0) + 1
        slug = base if counts[base] == 1 else f"{base}-{counts[base]}"
        slugs.append(slug)
        lines.append(f"{i}. [{doc.title}](#{slug})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Document bodies.
    for i, (doc, slug) in enumerate(zip(docs, slugs), 1):
        lines.append(f"## {i}. {doc.title}")
        lines.append("")
        lines.append(f"- **Source:** [{doc.url}]({doc.url})")
        if doc.ext:
            lines.append(f"- **Type:** {doc.ext.upper()}")
        if doc.error:
            lines.append(f"- **Status:** ⚠️ extraction failed — {doc.error}")
        for k, v in doc.meta.items():
            lines.append(f"- **{k.replace('_', ' ').title()}:** {v}")
        lines.append("")
        body = normalize_text(doc.text) if doc.text else ""
        if body:
            lines.append(body)
        elif not doc.error:
            lines.append("> _No extractable text. Use the source link above._")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--url", default=POLICIES_URL, help="Policies index page URL")
    parser.add_argument(
        "--out", default="leelanau_county_policies.md", help="Output Markdown file"
    )
    parser.add_argument(
        "--cache", default="./.policy_cache", help="Directory for downloaded files"
    )
    parser.add_argument(
        "--delay", type=float, default=1.0, help="Seconds to wait between downloads"
    )
    parser.add_argument(
        "--limit", type=int, default=0, help="Only process the first N documents (0 = all)"
    )
    parser.add_argument(
        "--expect", type=int, default=53, help="Expected document count (warning only)"
    )
    args = parser.parse_args(argv)

    session = make_session()

    print(f"Fetching index: {args.url}")
    try:
        index = fetch(session, args.url)
    except requests.HTTPError as exc:
        print(f"ERROR: could not fetch index page: {exc}", file=sys.stderr)
        return 2

    docs = extract_links(index.text, args.url)
    print(f"Found {len(docs)} document links.")
    if args.expect and len(docs) != args.expect:
        print(
            f"  NOTE: expected {args.expect}; got {len(docs)}. "
            "Inspect the page structure or adjust looks_like_document().",
            file=sys.stderr,
        )

    if args.limit:
        docs = docs[: args.limit]

    for i, doc in enumerate(docs, 1):
        print(f"[{i}/{len(docs)}] {doc.title}")
        try:
            download(session, doc, args.cache)
            extract_text(doc)
        except Exception as exc:  # noqa: BLE001
            doc.error = str(exc)
            print(f"    ! {exc}", file=sys.stderr)
        time.sleep(args.delay)

    markdown = build_markdown(docs, args.url)
    with open(args.out, "w", encoding="utf-8") as fh:
        fh.write(markdown)

    ok = sum(1 for d in docs if d.text and not d.error)
    print(f"\nWrote {args.out} ({len(markdown):,} chars).")
    print(f"  {ok}/{len(docs)} documents had extractable text.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
